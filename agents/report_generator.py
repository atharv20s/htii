"""
Report Generation Agent — Produces professional DOCX credit memos.

Takes research report + credit scorecard and generates a polished
Word document suitable for bank credit committees and hackathon judges.

Sections:
    1. Cover Page with company name, date, rating badge
    2. Executive Summary (LLM-generated)
    3. Company Overview
    4. Financial Analysis
    5. Credit Scorecard (visual table)
    6. Risk Assessment Matrix
    7. Red Flags & Concerns
    8. Lending Recommendation
    9. Disclaimer

Uses python-docx for Word document generation and Groq LLM for
intelligent summarization.
"""
import os
import sys
import re
import glob
import logging
from typing import Dict, Optional
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import OUTPUT_DIR, CHROMA_COLLECTION_NAME, CHROMA_DB_DIR

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


# ============================================================
# Styling Constants
# ============================================================

COLORS = {
    "primary": RGBColor(0x1A, 0x23, 0x7E),      # Deep navy blue
    "secondary": RGBColor(0x00, 0x96, 0xC7),     # Accent blue
    "success": RGBColor(0x2D, 0x6A, 0x4F),       # Green
    "warning": RGBColor(0xE9, 0x6C, 0x00),       # Orange
    "danger": RGBColor(0xD0, 0x00, 0x00),         # Red
    "dark": RGBColor(0x21, 0x21, 0x21),           # Near black
    "muted": RGBColor(0x6C, 0x75, 0x7D),          # Gray
    "light_bg": RGBColor(0xF8, 0xF9, 0xFA),       # Light background
    "white": RGBColor(0xFF, 0xFF, 0xFF),           # White
    "score_high": RGBColor(0x2D, 0x6A, 0x4F),     # Green for 70+
    "score_mid": RGBColor(0xE9, 0x6C, 0x00),      # Orange for 50-69
    "score_low": RGBColor(0xD0, 0x00, 0x00),       # Red for <50
}

RISK_COLORS = {
    "Very Low": COLORS["success"],
    "Low": COLORS["success"],
    "Moderate-Low": RGBColor(0x5F, 0xAD, 0x56),
    "Moderate": COLORS["warning"],
    "Moderate-High": RGBColor(0xE0, 0x4D, 0x01),
    "High": COLORS["danger"],
    "Very High": RGBColor(0xA0, 0x00, 0x00),
    "Default Risk": RGBColor(0x80, 0x00, 0x00),
}


def _score_color(score: int) -> RGBColor:
    """Get color based on score value."""
    if score >= 70:
        return COLORS["score_high"]
    elif score >= 50:
        return COLORS["score_mid"]
    else:
        return COLORS["score_low"]


def _set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_styled_paragraph(doc, text, style="Normal", bold=False, color=None,
                          size=None, alignment=None, space_after=None, space_before=None):
    """Add a paragraph with custom styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


# ============================================================
# Report Sections
# ============================================================

def _add_cover_page(doc, company_name: str, rating: str, risk_level: str, score: float):
    """Add a professional cover page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTELLI-CREDIT")
    run.font.size = Pt(36)
    run.font.color.rgb = COLORS["primary"]
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Credit Analysis Report")
    run.font.size = Pt(18)
    run.font.color.rgb = COLORS["muted"]

    doc.add_paragraph("")

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name.upper())
    run.font.size = Pt(28)
    run.font.color.rgb = COLORS["dark"]
    run.bold = True

    doc.add_paragraph("")

    # Rating badge
    rating_table = doc.add_table(rows=1, cols=3)
    rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Score cell
    cell = rating_table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Score: {score}/100")
    run.font.size = Pt(14)
    run.font.color.rgb = _score_color(int(score))
    run.bold = True

    # Rating cell
    cell = rating_table.cell(0, 1)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Rating: {rating}")
    run.font.size = Pt(14)
    run.font.color.rgb = COLORS["primary"]
    run.bold = True

    # Risk cell
    cell = rating_table.cell(0, 2)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Risk: {risk_level}")
    run.font.size = Pt(14)
    run.font.color.rgb = RISK_COLORS.get(risk_level, COLORS["muted"])
    run.bold = True

    # Date and footer
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS["muted"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by Intelli-Credit AI Platform")
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["muted"]
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Powered by LangGraph ReAct Agents • Groq LLM • Tavily AI")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["muted"]
    run.italic = True

    doc.add_page_break()


def _add_section_header(doc, title: str, subtitle: str = None):
    """Add a styled section header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.font.size = Pt(16)
    run.font.color.rgb = COLORS["primary"]
    run.bold = True

    # Underline bar
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("━" * 60)
    run.font.size = Pt(6)
    run.font.color.rgb = COLORS["secondary"]

    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS["muted"]
        run.italic = True


def _add_executive_summary(doc, company_name: str, research_text: str, scorecard_text: str):
    """Generate and add an LLM-powered executive summary."""
    _add_section_header(doc, "1. Executive Summary",
                        "AI-generated overview of credit analysis findings")

    # Generate summary using LLM
    try:
        from langchain_setup import get_llm
        llm = get_llm()

        prompt = f"""You are writing an EXECUTIVE SUMMARY for a bank credit committee report.
Company: {company_name}

Based on the research and scoring below, write a 4-5 paragraph executive summary covering:
1. Company introduction (1 sentence)
2. Key financial strengths
3. Key risks and concerns
4. Credit score and rating with brief justification
5. Lending recommendation

Keep it professional, concise, and factual. Use specific numbers where available.
Do NOT use headers or bullet points — write flowing paragraphs.

RESEARCH FINDINGS:
{research_text[:3000]}

CREDIT SCORECARD:
{scorecard_text[:2000]}

Write the executive summary now (4-5 paragraphs, no headers):"""

        response = llm.invoke(prompt)
        summary = response.content if hasattr(response, 'content') else str(response)
    except Exception as e:
        logger.warning(f"LLM summary failed: {e}")
        summary = (
            f"This report presents a comprehensive credit analysis of {company_name}, "
            f"covering financial health, credit ratings, regulatory compliance, litigation risks, "
            f"promoter background, sector positioning, and ESG factors. "
            f"The analysis was conducted using AI-powered research agents with data from "
            f"Indian financial databases, regulatory filings, and uploaded company documents."
        )

    # Add summary paragraphs
    for para in summary.strip().split("\n\n"):
        if para.strip():
            p = doc.add_paragraph()
            run = p.add_run(para.strip())
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS["dark"]
            p.paragraph_format.space_after = Pt(8)


def _add_research_section(doc, title: str, section_num: int, content: str):
    """Add a research findings section."""
    _add_section_header(doc, f"{section_num}. {title}")

    for para in content.strip().split("\n"):
        line = para.strip()
        if not line:
            continue
        if line.startswith("-") or line.startswith("•"):
            p = doc.add_paragraph(line[1:].strip(), style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)
        elif ":" in line and len(line.split(":")[0]) < 40:
            # Key-value pair
            p = doc.add_paragraph()
            key, value = line.split(":", 1)
            run_key = p.add_run(f"{key.strip()}: ")
            run_key.bold = True
            run_key.font.size = Pt(10)
            run_key.font.color.rgb = COLORS["dark"]
            run_val = p.add_run(value.strip())
            run_val.font.size = Pt(10)
            run_val.font.color.rgb = COLORS["dark"]
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = COLORS["dark"]


def _add_scorecard_table(doc, scorecard_text: str):
    """Add the credit scorecard as a styled table."""
    _add_section_header(doc, "Credit Scorecard",
                        "7-category weighted scoring model")

    # Parse scores from scorecard text
    categories = []
    lines = scorecard_text.split("\n")
    for line in lines:
        # Match lines like "  1. Financial Health          80      30%      24.0"
        match = re.match(r'\s*\d+\.\s+(.+?)\s{2,}(\d+)\s+(\d+%)\s+([\d.]+)', line)
        if match:
            categories.append({
                "name": match.group(1).strip(),
                "score": int(match.group(2)),
                "weight": match.group(3),
                "weighted": float(match.group(4)),
            })

    if not categories:
        _add_styled_paragraph(doc, "Scorecard data could not be parsed from the scoring report.",
                              color=COLORS["muted"], size=10)
        return

    # Create table
    table = doc.add_table(rows=len(categories) + 2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    headers = ["#", "Category", "Score", "Weight", "Weighted"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS["white"]
        _set_cell_shading(cell, "1A237E")

    # Data rows
    total_weighted = 0
    for idx, cat in enumerate(categories):
        row = table.rows[idx + 1]
        cells = row.cells
        values = [str(idx + 1), cat["name"], str(cat["score"]),
                  cat["weight"], f"{cat['weighted']:.1f}"]

        for i, val in enumerate(values):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(10)

            # Color the score cell
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.color.rgb = _score_color(cat["score"])
            elif i == 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Alternate row shading
        if idx % 2 == 0:
            for cell in cells:
                _set_cell_shading(cell, "F8F9FA")

        total_weighted += cat["weighted"]

    # Total row
    total_row = table.rows[-1]
    total_cells = total_row.cells
    total_cells[0].text = ""
    total_cells[1].text = ""
    p = total_cells[1].paragraphs[0]
    run = p.add_run("OVERALL SCORE")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["primary"]

    total_cells[2].text = ""
    total_cells[3].text = ""
    total_cells[4].text = ""
    p = total_cells[4].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{total_weighted:.1f}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _score_color(int(total_weighted))

    for cell in total_cells:
        _set_cell_shading(cell, "E8EAF6")


def _add_red_flags(doc, scorecard_text: str):
    """Extract and display red flags."""
    _add_section_header(doc, "Red Flags & Concerns",
                        "Critical items requiring attention")

    # Parse red flags from scorecard
    flags = []
    in_flags = False
    for line in scorecard_text.split("\n"):
        if "RED FLAGS" in line:
            in_flags = True
            continue
        if in_flags:
            if line.strip().startswith("-"):
                flags.append(line.strip()[1:].strip())
            elif line.strip() and not line.strip().startswith("-"):
                if any(kw in line for kw in ["KEY DATA", "LENDING", "═", "─"]):
                    break

    if not flags:
        _add_styled_paragraph(doc, "No significant red flags identified.",
                              color=COLORS["success"], size=11)
        return

    # Create flags table
    table = doc.add_table(rows=len(flags) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = ""
    p = header_cells[0].paragraphs[0]
    run = p.add_run("#")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["white"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_shading(header_cells[0], "D00000")

    header_cells[1].text = ""
    p = header_cells[1].paragraphs[0]
    run = p.add_run("Red Flag Description")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["white"]
    _set_cell_shading(header_cells[1], "D00000")

    # Flag rows
    for i, flag in enumerate(flags):
        row = table.rows[i + 1]
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"⚠ {i+1}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS["danger"]
        run.bold = True

        row.cells[1].text = ""
        p = row.cells[1].paragraphs[0]
        run = p.add_run(flag)
        run.font.size = Pt(10)

        if i % 2 == 0:
            _set_cell_shading(row.cells[0], "FFF3F3")
            _set_cell_shading(row.cells[1], "FFF3F3")


def _add_lending_recommendation(doc, scorecard_text: str, score: float, rating: str):
    """Add the final lending recommendation section."""
    _add_section_header(doc, "Lending Recommendation",
                        "Final credit decision based on comprehensive analysis")

    # Parse decision from scorecard
    decision = "CONDITIONAL APPROVE"
    conditions = ""
    confidence = "0.75"

    for line in scorecard_text.split("\n"):
        if "Decision:" in line:
            decision = line.split("Decision:")[-1].strip()
        elif "Conditions:" in line:
            conditions = line.split("Conditions:")[-1].strip()
        elif "Confidence:" in line and "LENDING" not in line:
            confidence = line.split("Confidence:")[-1].strip()

    # Decision box
    decision_table = doc.add_table(rows=1, cols=1)
    decision_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = decision_table.cell(0, 0)
    cell.text = ""

    if "APPROVE" in decision and "CONDITIONAL" not in decision and "REJECT" not in decision:
        bg_color = "2D6A4F"
        text_color = COLORS["white"]
        icon = "✅"
    elif "CONDITIONAL" in decision:
        bg_color = "E96C00"
        text_color = COLORS["white"]
        icon = "⚠️"
    else:
        bg_color = "D00000"
        text_color = COLORS["white"]
        icon = "❌"

    _set_cell_shading(cell, bg_color)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{icon}  {decision}  {icon}\n")
    run.font.size = Pt(24)
    run.font.color.rgb = text_color
    run.bold = True

    # Details
    doc.add_paragraph("")

    details_table = doc.add_table(rows=3, cols=2)
    details_table.style = "Table Grid"
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    detail_items = [
        ("Credit Score", f"{score}/100"),
        ("Credit Rating", rating),
        ("Confidence Level", confidence),
    ]

    for i, (key, val) in enumerate(detail_items):
        cells = details_table.rows[i].cells
        cells[0].text = ""
        p = cells[0].paragraphs[0]
        run = p.add_run(key)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS["primary"]
        _set_cell_shading(cells[0], "E8EAF6")

        cells[1].text = ""
        p = cells[1].paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(11)
        run.bold = True

    if conditions:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Conditions / Notes: ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS["primary"]
        run = p.add_run(conditions)
        run.font.size = Pt(11)


def _add_disclaimer(doc):
    """Add a disclaimer section."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DISCLAIMER")
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS["muted"]
    run.bold = True

    disclaimer_text = (
        "This credit analysis report has been generated by the Intelli-Credit AI Platform "
        "using automated research agents, machine learning models, and publicly available data. "
        "The findings, scores, and recommendations are indicative and should be used as a "
        "supplementary tool alongside traditional credit assessment processes. "
        "The AI system may not capture all relevant information, and the data sources "
        "may contain inaccuracies. This report does not constitute financial advice. "
        "All lending decisions should be validated by qualified credit officers and "
        "must comply with RBI guidelines and institutional credit policies."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(disclaimer_text)
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["muted"]
    run.italic = True


# ============================================================
# Parse Research Report Sections
# ============================================================

def _parse_research_sections(research_text: str) -> Dict[str, str]:
    """Parse the research report into named sections."""
    sections = {}
    current_section = None
    current_content = []

    section_markers = [
        "COMPANY OVERVIEW", "CREDIT RATING", "FINANCIAL ANALYSIS",
        "FINANCIAL OBSERVATIONS", "SECTOR BENCHMARK", "PROMOTER BACKGROUND",
        "REGULATORY RISK", "LITIGATION RISK", "IRREGULARITIES & ESG",
        "HYPOTHESIS TESTING", "FINAL CREDIT INSIGHT",
    ]

    for line in research_text.split("\n"):
        stripped = line.strip()
        # Check if this line is a section header
        matched = False
        for marker in section_markers:
            if marker in stripped.upper() and len(stripped) < len(marker) + 10:
                if current_section:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = marker
                current_content = []
                matched = True
                break

        if not matched and current_section:
            # Skip separator lines
            if not all(c in "─═━─" for c in stripped) and stripped:
                current_content.append(line)

    if current_section:
        sections[current_section] = "\n".join(current_content).strip()

    return sections


def _parse_scorecard_values(scorecard_text: str) -> Dict:
    """Extract key values from the scorecard text."""
    values = {
        "overall_score": 0.0,
        "credit_rating": "N/A",
        "risk_level": "N/A",
    }

    for line in scorecard_text.split("\n"):
        if "OVERALL SCORE:" in line:
            match = re.search(r'([\d.]+)/100', line)
            if match:
                values["overall_score"] = float(match.group(1))
        elif "CREDIT RATING:" in line:
            values["credit_rating"] = line.split("CREDIT RATING:")[-1].strip()
        elif "RISK LEVEL:" in line:
            values["risk_level"] = line.split("RISK LEVEL:")[-1].strip()

    return values


# ============================================================
# Main Report Generator
# ============================================================

def generate_report(
    company_name: str,
    research_text: str,
    scorecard_text: str,
    output_filename: str = None,
) -> str:
    """Generate a professional DOCX credit analysis report.

    Args:
        company_name: Name of the company
        research_text: Full research report text
        scorecard_text: Full credit scorecard text
        output_filename: Optional custom filename

    Returns:
        Path to the generated DOCX file
    """
    print(f"\n{'='*60}")
    print(f"  📄 REPORT GENERATION AGENT")
    print(f"  Company: {company_name}")
    print(f"  Format: Professional DOCX Credit Memo")
    print(f"{'='*60}")

    # Parse inputs
    print(f"\n  📋 Parsing research report and scorecard...")
    research_sections = _parse_research_sections(research_text)
    scorecard_values = _parse_scorecard_values(scorecard_text)

    score = scorecard_values["overall_score"]
    rating = scorecard_values["credit_rating"]
    risk = scorecard_values["risk_level"]

    print(f"  Score: {score}/100 | Rating: {rating} | Risk: {risk}")
    print(f"  Research sections found: {len(research_sections)}")

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = COLORS["dark"]

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Build the report ──
    print(f"\n  🔨 Building report sections...")

    # 1. Cover Page
    print(f"    ✓ Cover page")
    _add_cover_page(doc, company_name, rating, risk, score)

    # 2. Executive Summary (LLM-powered)
    print(f"    ✓ Executive summary (generating with LLM...)")
    _add_executive_summary(doc, company_name, research_text, scorecard_text)

    # 3. Research sections
    section_mapping = [
        ("COMPANY OVERVIEW", "Company Overview"),
        ("CREDIT RATING", "Credit Rating Analysis"),
        ("FINANCIAL ANALYSIS", "Financial Analysis"),
        ("FINANCIAL OBSERVATIONS", "Financial Analysis"),
        ("SECTOR BENCHMARK", "Sector & Market Position"),
        ("PROMOTER BACKGROUND", "Promoter Background"),
        ("REGULATORY RISK", "Regulatory Risk Assessment"),
        ("LITIGATION RISK", "Litigation Risk Assessment"),
        ("IRREGULARITIES & ESG", "ESG & Irregularities"),
        ("HYPOTHESIS TESTING", "Hypothesis Testing Results"),
        ("FINAL CREDIT INSIGHT", "Credit Insight Summary"),
    ]

    section_num = 2
    added_sections = set()
    for marker, display_name in section_mapping:
        if marker in research_sections and display_name not in added_sections:
            section_num += 1
            print(f"    ✓ {display_name}")
            _add_research_section(doc, display_name, section_num,
                                  research_sections[marker])
            added_sections.add(display_name)

    # Credit Scorecard Table
    section_num += 1
    print(f"    ✓ Credit scorecard table")
    _add_scorecard_table(doc, scorecard_text)

    # Red Flags
    print(f"    ✓ Red flags")
    _add_red_flags(doc, scorecard_text)

    # Lending Recommendation
    print(f"    ✓ Lending recommendation")
    doc.add_page_break()
    _add_lending_recommendation(doc, scorecard_text, score, rating)

    # Disclaimer
    print(f"    ✓ Disclaimer")
    _add_disclaimer(doc)

    # ── Save ──
    if not output_filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = company_name.replace(" ", "_").lower()
        output_filename = f"credit_report_{safe_name}_{timestamp}.docx"

    filepath = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(filepath)

    print(f"\n  {'='*55}")
    print(f"  ✅ Report saved to: {filepath}")
    print(f"  📄 File size: {os.path.getsize(filepath) / 1024:.1f} KB")
    print(f"  {'='*55}\n")

    return filepath


# ============================================================
# Standalone Execution
# ============================================================

if __name__ == "__main__":
    print("\n  ╔══════════════════════════════════════════════╗")
    print("  ║  Intelli-Credit Report Generator             ║")
    print("  ╚══════════════════════════════════════════════╝\n")

    # Find research reports
    research_files = glob.glob(os.path.join(OUTPUT_DIR, "research_*.txt")) + \
                     glob.glob(os.path.join(OUTPUT_DIR, "test_research_*.txt"))

    # Find scorecards
    scorecard_files = glob.glob(os.path.join(OUTPUT_DIR, "scorecard_*.txt"))

    if not research_files:
        print("  ❌ No research reports found in data/output/")
        print("  Run the research agent first.")
        sys.exit(1)

    if not scorecard_files:
        print("  ❌ No scorecards found in data/output/")
        print("  Run the credit scoring agent first.")
        sys.exit(1)

    # Select research report
    print("  Research Reports:")
    for i, f in enumerate(research_files, 1):
        print(f"    [{i}] {os.path.basename(f)}")
    r_choice = input("\n  Select research report > ").strip()

    try:
        r_idx = int(r_choice) - 1
        research_path = research_files[r_idx]
    except (ValueError, IndexError):
        print("  Invalid choice.")
        sys.exit(1)

    # Select scorecard
    print("\n  Scorecards:")
    for i, f in enumerate(scorecard_files, 1):
        print(f"    [{i}] {os.path.basename(f)}")
    s_choice = input("\n  Select scorecard > ").strip()

    try:
        s_idx = int(s_choice) - 1
        scorecard_path = scorecard_files[s_idx]
    except (ValueError, IndexError):
        print("  Invalid choice.")
        sys.exit(1)

    # Read files
    with open(research_path, "r", encoding="utf-8") as f:
        research_text = f.read()
    with open(scorecard_path, "r", encoding="utf-8") as f:
        scorecard_text = f.read()

    # Extract company name from filename
    basename = os.path.basename(research_path)
    name_part = basename.replace("test_research_", "").replace("research_", "")
    name_part = name_part.rsplit("_", 1)[0] if "_20" in name_part else name_part.replace(".txt", "")
    company_name = name_part.replace("_", " ").title()

    confirm = input(f"\n  Company: {company_name}. Correct? (y/n) > ").strip()
    if confirm.lower() != "y":
        company_name = input("  Enter correct company name > ").strip()

    # Generate report
    generate_report(company_name, research_text, scorecard_text)
